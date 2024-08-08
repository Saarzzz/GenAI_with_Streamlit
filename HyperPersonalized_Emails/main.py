import os
import time
from flask import Flask, request, jsonify, render_template, redirect, url_for, send_from_directory
from werkzeug.utils import secure_filename
from flask_cors import CORS
from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image
from langchain_core.documents import Document  # Ensure using the correct Document class
import pandas as pd  # Import pandas for reading Excel files

# Loading environment variables
load_dotenv()

# Loading the GROQ and Gemini API keys
groq_api_key = os.getenv('GROQ_API_KEY')
gemini_api_key = os.getenv('GEMINI_API_KEY')

# Initialize Flask app
app = Flask(__name__)
CORS(app)
app.config['UPLOAD_FOLDER'] = 'uploads'
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# Initialize LLM and prompt template
llm = ChatGroq(groq_api_key=groq_api_key, model_name="Llama3-8b-8192", temperature=0.1)
prompt = ChatPromptTemplate.from_template(
    """
    Answer the questions based on the provided context only.
    Please provide the most accurate response based on the question
    <context>
    {context}
    <context>
    Questions: {input}
    Always follow these guidelines:
    1. Provide a brief description without using phrases like "Based on the provided context."
    2. Be as detailed as possible.
    3. Always use the <br> tag for line breaks instead of '\n' in the output.
    4. When providing a link, use the <a href='link' target='_blank'> tag for redirection.
    5. Do not use information other than the provided context.
    """
)

# Helper functions to read different file types
def read_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
    return text

def read_docx(file_path):
    doc = DocxDocument(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text

def read_image(file_path):
    return file_path.replace("\\", "/"), os.path.basename(file_path)  # Return the file path and file name

def read_xlsx(file_path):
    df = pd.read_excel(file_path)
    rows = []
    for index, row in df.iterrows():
        row_str = f"{index + 1}." + " ".join([f"{col}={row[col]}," for col in df.columns])
        rows.append(row_str)
    text = "\n".join(rows)
    return text

def load_documents(folder_path):
    documents = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith('.pdf'):
            text = read_pdf(file_path)
            documents.append(Document(page_content=text, metadata={"source": filename}))
        elif filename.endswith('.docx'):
            text = read_docx(file_path)
            documents.append(Document(page_content=text, metadata={"source": filename}))
        elif filename.endswith('.xlsx'):
            text = read_xlsx(file_path)
            documents.append(Document(page_content=text, metadata={"source": filename}))
        elif filename.endswith('.png'):
            image_path, image_name = read_image(file_path)
            documents.append(Document(page_content=image_path, metadata={"source": image_name}))
    return documents

def vector_embedding():
    start_time = time.time()
    app.config["embeddings"] = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    folder_path = app.config['UPLOAD_FOLDER']
    app.config["docs"] = load_documents(folder_path)  # Document Loading
    print(f"Document loading time: {time.time() - start_time:.2f} seconds")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=100)  # Chunk Creation
    app.config["final_documents"] = text_splitter.split_documents(app.config["docs"])  # Splitting
    print(f"Document splitting time: {time.time() - start_time:.2f} seconds")
    app.config["vectors"] = FAISS.from_documents(app.config["final_documents"], app.config["embeddings"])  # Vector embeddings
    print(f"Embedding generation time: {time.time() - start_time:.2f} seconds")
    # Debug: Print the content of the documents being embedded
    #for doc in app.config["final_documents"]:
        #print("Embedding document:", doc.page_content)
    print(f"Total embedding task time: {time.time() - start_time:.2f} seconds")

# Route to render the main page
@app.route('/')
def index():
    return render_template('index.html')

# Route to upload documents
@app.route('/upload_documents', methods=['POST'])
def upload_documents():
    if 'files[]' not in request.files:
        return redirect(request.url)
    files = request.files.getlist('files[]')
    for file in files:
        if file:
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            try:
                file.save(file_path)
                print(f"File saved to {file_path}")
            except Exception as e:
                print(f"Error saving file {filename}: {e}")
                return jsonify({"error": f"Failed to save file {filename}: {e}"}), 500
    return jsonify({"status": "Documents uploaded successfully"})

# Route to process documents
@app.route('/process_documents', methods=['POST'])
def process_documents():
    vector_embedding()
    return jsonify({"status": "Documents processed successfully"})

# Route to serve uploaded files
@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# Route to process queries
@app.route('/query', methods=['POST'])
def query():
    data = request.json
    input_query = data.get('input')
    if not input_query:
        return jsonify({"error": "No input query provided"}), 400

    document_chain = create_stuff_documents_chain(llm, prompt)
    retriever = app.config["vectors"].as_retriever()
    retrieval_chain = create_retrieval_chain(retriever, document_chain)
    start = time.process_time()
    response = retrieval_chain.invoke({'input': input_query})
    print("Response time:", time.process_time() - start)
    #print(response['answer'])
    result = []
    displayed_image = False
    for doc in response["context"]:
        if doc.metadata["source"].endswith('.png') and not displayed_image:
            result.append({"type": "image", "content": doc.page_content, "caption": doc.metadata["source"]})
            displayed_image = True
    #print(result)
    return jsonify({"answer": response['answer'], "result" : result })

# Run the Flask app
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)

