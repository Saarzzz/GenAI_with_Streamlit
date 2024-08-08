import os
import streamlit as st
from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_core.documents import Document  
import pandas as pd

# Load environment variables
load_dotenv()

# Load the GROQ and Gemini API keys
groq_api_key = os.getenv('GROQ_API_KEY')
gemini_api_key = os.getenv('GEMINI_API_KEY')

st.title("Real Estate QnA System")

llm = ChatGroq(groq_api_key=groq_api_key, model_name="Llama3-8b-8192", temperature=0.4)

prompt = ChatPromptTemplate.from_template(
"""
    Answer the questions based on the provided context only.
    Please provide the most accurate response based on the question
    <context>
    {context}
    <context>
    Questions: {input}
    Always follow these guidelines:
    1. Provide a brief description without using phrases like "Based on the provided context."
    2. Be as detailed as possible.
    3. Always use the <br> tag for line breaks instead of '\n' in the output.
    4. Do not use information other than the provided context.
    """
)

def read_pdf(file):  # Reading the PDF Files
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
    return text

def read_docx(file):  # Reading the DOCX Files
    doc = DocxDocument(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text

def read_image(file):  # Reading the Image Files
    image = Image.open(file)
    return image
    
def read_xlsx(file):  # Reading the XLSX Files
    df = pd.read_excel(file)
    text = df.to_string(index=False)
    return text

def load_documents(files):
    documents = []
    for file in files:
        if file.name.endswith('.pdf'):
            text = read_pdf(file)
            documents.append(Document(page_content=text, metadata={"source": file.name}))
        elif file.name.endswith('.docx'):
            text = read_docx(file)
            documents.append(Document(page_content=text, metadata={"source": file.name}))
        elif file.name.endswith('.xlsx'):
            text = read_xlsx(file)
            documents.append(Document(page_content=text, metadata={"source": file.name}))
        elif file.name.endswith('.png'):
            image = read_image(file)
            documents.append(Document(page_content=image, metadata={"source": file.name}))
    return documents

def vector_embedding(files):
    if "vectors" not in st.session_state:
        st.session_state.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        st.session_state.docs = load_documents(files)  # Document Loading
        st.session_state.text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=100)  # Chunk Creation
        st.session_state.final_documents = st.session_state.text_splitter.split_documents(st.session_state.docs)  # Splitting
        st.session_state.vectors = FAISS.from_documents(st.session_state.final_documents, st.session_state.embeddings)  # Vector embeddings

# Streamlit UI
uploaded_files = st.file_uploader("Upload your documents", type=['pdf', 'docx', 'xlsx', 'png'], accept_multiple_files=True)

if st.button("Load the Embeddings") and uploaded_files:
    vector_embedding(uploaded_files)
    st.write("Vector Store DB Is Ready")
    
prompt1 = st.text_input("Enter Your Question:-")

import time

if prompt1:
    document_chain = create_stuff_documents_chain(llm, prompt)
    retriever = st.session_state.vectors.as_retriever()
    retrieval_chain = create_retrieval_chain(retriever, document_chain)
    start = time.process_time()
    response = retrieval_chain.invoke({'input': prompt1})
    print("Response time:", time.process_time() - start)
    st.write(response['answer'])
    displayed_image = False
    for doc in response["context"]:
    # Check if the document source ends with '.png'
        if doc.metadata["source"].endswith('.png') and not displayed_image:
            # Ensure that the page_content is an image object
            if isinstance(doc.page_content, Image.Image):
                st.image(doc.page_content, caption=doc.metadata["source"])
                displayed_image = True
            else:
                st.error(f"Error: The content for {doc.metadata['source']} is not a valid image.")
         
    # With a Streamlit expander
    with st.expander("Document Similarity Search"):
        # Find the relevant chunks
        for i, doc in enumerate(response["context"]):
            st.write(doc.page_content)
            st.write("--------------------------------")

      # Display images based on file names in the response
        for doc in response["context"]:
          if doc.metadata["source"].endswith('.png'):
            st.image(doc.page_content, caption=doc.metadata["source"])
